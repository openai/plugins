from __future__ import annotations

import tempfile
import unittest
from pathlib import Path

from bs4 import BeautifulSoup
from docx import Document

from report_to_google_doc.docx_writer import write_docx
from report_to_google_doc.html_parser import parse_html, portable_chart_source_note
from report_to_google_doc.quality import build_preflight_checks


PORTABLE_REPORT = """<!doctype html>
<html>
<head><meta charset="utf-8"><title>Portable conversion fixture</title></head>
<body>
<main id="data-analytics-portable-fallback" class="portable-fallback">
  <header><h1>Portable conversion fixture</h1></header>
  <section class="portable-markdown">
    <h2>Conclusions</h2>
    <p>Conversion-adjusted pipeline reached <strong><span class="portable-source-tooltip portable-source-value" data-portable-source-host="true" tabindex="0" aria-describedby="portable-source-tooltip-narrative-value"><span class="portable-source-value-text">$6.73M</span><span class="portable-source-tooltip-content" id="portable-source-tooltip-narrative-value" role="tooltip"><span class="portable-source-tooltip-heading" aria-hidden="true">Source for $6.73M</span><strong>Source: Narrative-only tooltip</strong><span class="portable-source-meta">Table: inline.narrative_only</span></span></span></strong> and delivered a <span class="portable-source-tooltip portable-source-value" data-portable-source-host="true" tabindex="0" aria-describedby="portable-source-tooltip-narrative-lift"><span class="portable-source-value-text">12.8%</span><span class="portable-source-tooltip-content" id="portable-source-tooltip-narrative-lift" role="tooltip"><span class="portable-source-tooltip-heading" aria-hidden="true">Source for 12.8%</span><strong>Source: Narrative-only tooltip</strong><span class="portable-source-meta">Table: inline.narrative_only</span></span></span> lift.</p>
  </section>
  <section class="portable-metric-grid">
    <article class="portable-metric-card">
      <p class="portable-metric-label">Revenue</p>
      <p class="portable-metric-value">
        <span class="portable-source-value" data-portable-source-host="true" tabindex="0">
          <span class="portable-source-value-text">$7.42M</span>
          <span class="portable-source-tooltip-content" role="tooltip">
            <strong>Source: Inline metric tooltip</strong>
            <span class="portable-source-meta">Table: inline.metric_only</span>
          </span>
        </span>
      </p>
      <p class="portable-card-description">Current quarter revenue and growth.</p>
      <div class="portable-metric-badges">
        <span class="portable-metric-badge"><span>Growth</span> <strong><span class="portable-source-value" data-portable-source-host="true" tabindex="0"><span class="portable-source-value-text">+18.4%</span><span class="portable-source-tooltip-content" role="tooltip"><strong>Source: Inline badge tooltip</strong></span></span></strong></span>
      </div>
      <div class="portable-inline-source portable-source-summary" data-source-id="revenue_sql">
        <div class="portable-inline-source-content portable-source-summary-content">
          <strong>Source: Revenue warehouse query</strong>
          <span class="portable-source-meta">Table: analytics.revenue</span>
          <pre class="portable-source-query-data"><code>SELECT revenue FROM analytics.revenue</code></pre>
        </div>
      </div>
    </article>
    <article class="portable-metric-card">
      <p class="portable-metric-label">Gross margin</p>
      <p class="portable-metric-value">71.3%</p>
      <p class="portable-card-description">Current gross margin.</p>
    </article>
  </section>
  <figure class="portable-content-card portable-chart-summary" data-chart-id="quarterly_revenue" data-portable-visual-title="Revenue is accelerating" data-portable-source-host="true" tabindex="0" aria-label="Revenue is accelerating" aria-describedby="portable-source-tooltip-1">
    <figcaption class="portable-visual-header portable-markdown">
      <p>Revenue is <strong>accelerating</strong></p>
      <p>Actual revenue has remained above plan.</p>
    </figcaption>
    <div class="portable-static-chart" role="img" aria-label="Revenue is accelerating chart">
      <div class="portable-static-chart-variant portable-static-chart-light" aria-hidden="true">
        <svg class="portable-static-chart-svg" aria-hidden="true" focusable="false" viewBox="0 0 640 360" width="640" height="360">
          <line x1="64" y1="300" x2="600" y2="300" stroke="#d0d0d0"></line>
          <rect x="120" y="120" width="120" height="180" fill="#1473e6"></rect>
          <rect x="320" y="80" width="120" height="220" fill="#6f6e69"></rect>
          <path d="M80 260 L220 190 L380 140 L560 90" fill="none" stroke="#1473e6"></path>
        </svg>
      </div>
      <div class="portable-static-chart-variant portable-static-chart-dark" aria-hidden="true">
        <svg class="portable-static-chart-svg" aria-hidden="true" focusable="false" viewBox="0 0 640 360" width="640" height="360">
          <line x1="64" y1="300" x2="600" y2="300" stroke="#505050"></line>
          <rect x="120" y="120" width="120" height="180" fill="#66a8ff"></rect>
          <rect x="320" y="80" width="120" height="220" fill="#aaa9a4"></rect>
          <path d="M80 260 L220 190 L380 140 L560 90" fill="none" stroke="#66a8ff"></path>
        </svg>
      </div>
    </div>
    <div class="portable-chart-data portable-chart-data-has-vector">
      <div class="portable-table-scroll">
        <table>
          <caption>Revenue is accelerating data</caption>
          <thead><tr><th>Quarter</th><th>Revenue</th><th>Series</th></tr></thead>
          <tbody>
            <tr><td>Q1</td><td><span class="portable-source-value" data-portable-source-host="true" tabindex="0"><span class="portable-source-value-text">$5.8M</span><span class="portable-inline-source-content portable-source-tooltip-content" role="tooltip"><strong>Source: Inline table tooltip</strong><span class="portable-source-meta">Table: inline.table_only</span></span></span></td><td>Actual</td></tr>
            <tr><td>Q1</td><td>$5.6M</td><td>Plan</td></tr>
          </tbody>
        </table>
      </div>
    </div>
    <div class="portable-inline-source portable-source-summary" data-source-id="revenue_sql">
      <div class="portable-inline-source-content portable-source-summary-content">
        <span class="portable-source-tooltip-heading" aria-hidden="true">Source for Revenue is accelerating</span>
        <strong>Source: Revenue warehouse query</strong>
        <span class="portable-source-meta">Table: analytics.revenue</span>
        <p class="portable-source-description-data">Reviewed quarterly revenue rows.</p>
        <pre class="portable-source-query-data"><code>SELECT quarter, revenue FROM analytics.revenue</code></pre>
      </div>
    </div>
  </figure>
</main>
</body>
</html>
"""


class PortableFallbackConversionTest(unittest.TestCase):
    def setUp(self) -> None:
        self.temporary_directory = tempfile.TemporaryDirectory()
        self.root = Path(self.temporary_directory.name)
        self.html = self.root / "report.html"
        self.html.write_text(PORTABLE_REPORT, encoding="utf-8")

    def tearDown(self) -> None:
        self.temporary_directory.cleanup()

    def test_parser_preserves_metric_grid_and_labeled_chart_data(self) -> None:
        manifest = parse_html(self.html)

        metric_table = next(table for table in manifest["tables"] if table["kind"] == "metric_cards")
        chart_table = next(table for table in manifest["tables"] if table["kind"] == "chart_data")

        self.assertEqual(manifest["counts"]["portable_metric_source_cards"], 2)
        self.assertEqual(manifest["counts"]["portable_metric_table_cells"], 2)
        self.assertEqual(metric_table["source_kind"], "portable_metric_grid")
        self.assertEqual(
            metric_table["rows"][0][0],
            "$7.42M\nRevenue\nCurrent quarter revenue and growth.\nGrowth +18.4%",
        )

        self.assertEqual(manifest["counts"]["portable_chart_source_blocks"], 1)
        self.assertEqual(manifest["counts"]["portable_chart_data_tables"], 1)
        self.assertEqual(manifest["counts"]["images"], 0)
        self.assertEqual(manifest["counts"]["svgs"], 2)
        self.assertEqual(manifest["counts"]["portable_static_chart_svgs"], 2)
        self.assertEqual(manifest["counts"]["chart_images"], 0)
        self.assertEqual(manifest["counts"]["svg_image_blocks"], 0)
        self.assertEqual(chart_table["semantic_label"], "Revenue is accelerating")
        self.assertEqual(chart_table["rows"][0], ["Quarter", "Revenue", "Series"])
        self.assertEqual(chart_table["rows"][1], ["Q1", "$5.8M", "Actual"])
        self.assertNotIn("Inline metric tooltip", metric_table["rows"][0][0])
        self.assertNotIn("Inline badge tooltip", metric_table["rows"][0][0])
        self.assertNotIn("Inline table tooltip", " ".join(chart_table["rows"][1]))
        self.assertIn("Revenue is accelerating\n", manifest["skeleton_text"])
        self.assertIn("Actual revenue has remained above plan.", manifest["skeleton_text"])
        self.assertIn(
            "Conversion-adjusted pipeline reached $6.73M and delivered a 12.8% lift.",
            manifest["skeleton_text"],
        )
        self.assertEqual(manifest["skeleton_text"].count("$6.73M"), 1)
        self.assertEqual(manifest["skeleton_text"].count("12.8%"), 1)
        self.assertNotIn("Narrative-only tooltip", manifest["skeleton_text"])
        self.assertNotIn("inline.narrative_only", manifest["skeleton_text"])
        self.assertIn("Source: Revenue warehouse query", manifest["skeleton_text"])
        self.assertNotIn("inline.table_only", manifest["skeleton_text"])
        self.assertFalse(any(table["kind"] == "table" for table in manifest["tables"]))

    def test_preflight_and_docx_keep_portable_structures(self) -> None:
        manifest = parse_html(self.html)
        preflight = build_preflight_checks(manifest, self.root)

        self.assertEqual(preflight["status"], "passed")
        portable_checks = {
            check["name"]: check["status"]
            for check in preflight["checks"]
            if check["name"].startswith("portable_")
        }
        self.assertEqual(
            portable_checks,
            {
                "portable_chart_data_is_labeled": "passed",
                "portable_chart_summaries_preserved": "passed",
                "portable_metric_cards_preserved": "passed",
            },
        )

        docx_path = write_docx(manifest, self.root)
        document = Document(docx_path)
        self.assertEqual(len(document.tables), 2)
        self.assertEqual(document.tables[0].cell(0, 0).text.splitlines()[0], "$7.42M")
        metric_runs = [
            run
            for run in document.tables[0].cell(0, 0).paragraphs[0].runs
            if run.text.strip()
        ]
        self.assertTrue(metric_runs[0].bold)
        self.assertEqual(document.tables[1].cell(0, 0).text, "Quarter")
        chart_header_runs = [
            run
            for run in document.tables[1].cell(0, 0).paragraphs[0].runs
            if run.text.strip()
        ]
        self.assertTrue(chart_header_runs[0].bold)
        self.assertIn(
            "Revenue is accelerating",
            [paragraph.text for paragraph in document.paragraphs],
        )
        document_text = "\n".join(paragraph.text for paragraph in document.paragraphs)
        self.assertEqual(document_text.count("$6.73M"), 1)
        self.assertEqual(document_text.count("12.8%"), 1)
        self.assertNotIn("Narrative-only tooltip", document_text)
        self.assertNotIn("inline.narrative_only", document_text)

    def test_legacy_source_tooltip_remains_a_provenance_fallback(self) -> None:
        soup = BeautifulSoup(
            """
            <figure class="portable-content-card portable-chart-summary">
              <div class="portable-inline-source">
                <div class="portable-inline-source-content portable-source-tooltip-content">
                  <strong>Source: Legacy warehouse query</strong>
                  <span class="portable-source-meta">Table: analytics.legacy</span>
                </div>
              </div>
            </figure>
            """,
            "html.parser",
        )
        figure = soup.find("figure")
        self.assertIsNotNone(figure)
        self.assertEqual(
            portable_chart_source_note(figure),
            "Source: Legacy warehouse query · Table: analytics.legacy",
        )


if __name__ == "__main__":
    unittest.main()
